from fastapi import FastAPI, UploadFile, File
from fastapi.responses import FileResponse
from fastapi.staticfiles import StaticFiles
import os
import uuid

import pdfplumber
from pdf2image import convert_from_path
import pytesseract
import cv2
import numpy as np
from docx import Document

app = FastAPI()

# Folders
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
UPLOADS_DIR = os.path.join(BASE_DIR, "uploads")
STATIC_DIR = os.path.join(BASE_DIR, "..", "static")

os.makedirs(UPLOADS_DIR, exist_ok=True)

# Serve frontend
app.mount("/", StaticFiles(directory=STATIC_DIR, html=True), name="static")


@app.post("/convert")
async def convert(file: UploadFile = File(...)):
    pdf_path = os.path.join(UPLOADS_DIR, f"{uuid.uuid4()}.pdf")

    with open(pdf_path, "wb") as f:
        f.write(await file.read())

    doc = Document()
    text_found = False

    # 1️⃣ Try extracting normal text
    with pdfplumber.open(pdf_path) as pdf:
        for page in pdf.pages:
            text = page.extract_text()
            if text and text.strip():
                doc.add_paragraph(text)
                text_found = True

    # 2️⃣ OCR fallback for scanned PDFs
    if not text_found:
        images = convert_from_path(pdf_path, dpi=300)
        for img in images:
            gray = cv2.cvtColor(np.array(img), cv2.COLOR_BGR2GRAY)
            gray = cv2.threshold(
                gray, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU
            )[1]

            text = pytesseract.image_to_string(
                gray,
                lang="eng",
                config="--oem 3 --psm 6"
            )

            if text.strip():
                doc.add_paragraph(text)

    output_path = os.path.join(UPLOADS_DIR, f"{uuid.uuid4()}.docx")
    doc.save(output_path)

    return FileResponse(
        output_path,
        filename="output.docx",
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
